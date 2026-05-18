#!/usr/bin/env python3
"""docx_toolkit — 给任意 docx 文件添加 Word 原生的"自动目录字段"和"中文页脚页码"。

用法:
    python docx_toolkit.py <input.docx> [-o output.docx] [--toc] [--pagenum]
"""

import argparse
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
TOC_PLACEHOLDER = "[请在 Word 中按 Ctrl+A 全选后按 F9 更新目录]"


def _make_field_run(instr):
    """构造一个 simple field run: <w:r><w:fldChar begin/><w:instrText/><w:fldChar end/></w:r>。"""
    r = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(XML_SPACE, "preserve")
    it.text = instr
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r.extend([fc1, it, fc2])
    return r


def _make_text_run(text):
    """构造一个静态文本 run: <w:r><w:t xml:space='preserve'>text</w:t></w:r>。"""
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    return r


def _add_page_number_footer(section):
    """在 section 页脚插入居中的中文页码: "第 X 页 / 共 Y 页"。"""
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = para._p
    p.append(_make_text_run("第 "))
    p.append(_make_field_run(" PAGE "))
    p.append(_make_text_run(" 页 / 共 "))
    p.append(_make_field_run(" NUMPAGES "))
    p.append(_make_text_run(" 页"))


def _insert_toc_field(doc):
    """在文档最开头插入: 目录(Heading1) → TOC字段(带占位) → 分页符。"""
    body = doc.element.body

    # TOC 字段段落 (begin → instrText → separate → 占位文本 → end)
    p_toc = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(XML_SPACE, "preserve")
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    fc_s = OxmlElement("w:fldChar")
    fc_s.set(qn("w:fldCharType"), "separate")
    r1.extend([fc1, it, fc_s])

    r_ph = OxmlElement("w:r")
    t_ph = OxmlElement("w:t")
    t_ph.text = TOC_PLACEHOLDER
    r_ph.append(t_ph)

    r2 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r2.append(fc2)

    p_toc.extend([r1, r_ph, r2])

    # "目录" 标题 (Heading1)
    p_h = OxmlElement("w:p")
    pPr_h = OxmlElement("w:pPr")
    pSt = OxmlElement("w:pStyle")
    pSt.set(qn("w:val"), "Heading1")
    pPr_h.append(pSt)
    p_h.append(pPr_h)
    r_h = OxmlElement("w:r")
    t_h = OxmlElement("w:t")
    t_h.text = "目录"
    r_h.append(t_h)
    p_h.append(r_h)

    # 分页符
    p_br = OxmlElement("w:p")
    r_br = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r_br.append(br)
    p_br.append(r_br)

    # 倒序 insert 到 body 开头, 顺序为: 目录 → TOC → 分页符
    body.insert(0, p_br)
    body.insert(0, p_toc)
    body.insert(0, p_h)


def main():
    parser = argparse.ArgumentParser(
        description="给 docx 加自动目录字段和中文页脚页码。",
    )
    parser.add_argument("input", help="输入的 docx 文件路径")
    parser.add_argument("-o", "--output", help="输出文件路径 (默认: <input>_processed.docx)")
    parser.add_argument("--toc", action="store_true", help="在文档开头插入自动目录字段")
    parser.add_argument("--pagenum", action="store_true", help="在所有 section 页脚插入中文页码")
    args = parser.parse_args()

    if not (args.toc or args.pagenum):
        parser.error("至少需要指定 --toc 或 --pagenum 之一")

    input_path = os.path.abspath(args.input)
    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        base, ext = os.path.splitext(input_path)
        output_path = base + "_processed" + (ext or ".docx")

    try:
        doc = Document(input_path)

        if args.toc:
            _insert_toc_field(doc)

        if args.pagenum:
            for section in doc.sections:
                _add_page_number_footer(section)

        doc.save(output_path)
    except Exception as e:
        print(f"[ERROR] 处理失败: {e}", file=sys.stderr)
        sys.exit(1)

    print(output_path)


if __name__ == "__main__":
    main()
